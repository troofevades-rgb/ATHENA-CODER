"""Document fixtures — synthesised at test time (T4-05).

Same reproducibility story as the rest of athena's media
fixtures: build deterministic PDFs / DOCXs on demand instead
of checking binary blobs into the repo.

PDF fixtures use PyMuPDF's writer; DOCX fixtures use
python-docx. Both libraries are required for the extractor
tests anyway — if they're not installed, the consumer tests
skip cleanly.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

import pytest


def make_pdf_with_text(
    out: Path,
    *,
    pages: list[str] | None = None,
    title: str = "Test PDF",
    author: str = "athena",
) -> Path:
    """Synthesise a PDF with text on each page. Default: 2
    pages of distinct text so reading-order tests can verify
    page order is preserved."""
    import fitz

    pages = pages or [
        "First page heading\nFirst page body paragraph one.\nFirst page body paragraph two.",
        "Second page heading\nSecond page body paragraph one.\nSecond page body paragraph two.",
    ]
    doc = fitz.open()
    for body in pages:
        page = doc.new_page()
        page.insert_text((72, 72), body, fontsize=12)
    doc.set_metadata({"title": title, "author": author})
    doc.save(str(out))
    doc.close()
    return out


def make_pdf_with_outline(out: Path) -> Path:
    """PDF with a programmatic outline (table of contents).
    Used to test outline extraction."""
    import fitz

    doc = fitz.open()
    titles = ["Introduction", "Methods", "Results", "Conclusion"]
    for t in titles:
        page = doc.new_page()
        page.insert_text((72, 72), t + "\n\nBody text for " + t, fontsize=12)
    # set_toc takes [[level, title, page_no_1indexed], ...]
    doc.set_toc([
        [1, "Introduction", 1],
        [1, "Methods", 2],
        [2, "Sub-method A", 2],
        [1, "Results", 3],
        [1, "Conclusion", 4],
    ])
    doc.save(str(out))
    doc.close()
    return out


def make_pdf_with_scanned_page(out: Path) -> Path:
    """A 2-page PDF where page 1 has text and page 2 is blank
    (no text layer — simulates a scanned page). The tool's
    OCR-fallback path keys off the empty-text-layer signal."""
    import fitz

    doc = fitz.open()
    p1 = doc.new_page()
    p1.insert_text((72, 72), "This page has real text.\nLine two.\nLine three.", fontsize=12)
    doc.new_page()  # empty — no text → flagged as scanned
    doc.save(str(out))
    doc.close()
    return out


def make_docx_with_structure(out: Path) -> Path:
    """DOCX with Title + multiple Heading-N styles + body
    paragraphs + a table. Used to test outline + structure +
    table extraction."""
    from docx import Document

    d = Document()
    d.core_properties.title = "Test DOCX"
    d.core_properties.author = "athena"
    d.add_heading("Test DOCX", level=0)  # Title style
    d.add_heading("Section A", level=1)
    d.add_paragraph("Paragraph one in section A.")
    d.add_heading("Sub-section A1", level=2)
    d.add_paragraph("Paragraph in sub-section.")
    d.add_heading("Section B", level=1)
    d.add_paragraph("Paragraph in section B.")
    # Add a 2x3 table
    tbl = d.add_table(rows=2, cols=3)
    headers = tbl.rows[0].cells
    headers[0].text = "Name"
    headers[1].text = "Age"
    headers[2].text = "Role"
    row = tbl.rows[1].cells
    row[0].text = "Alice"
    row[1].text = "30"
    row[2].text = "Engineer"
    d.save(str(out))
    return out


# ---------------------------------------------------------------
# Stub OCR backend — for the OCR-fallback tests in T4-05.2
# ---------------------------------------------------------------


@pytest.fixture
def stub_ocr_recognize():
    """A stand-in ocr_recognize-compatible callable that
    returns a canned transcription per call. The tool layer
    injects this via the _ocr_fn argument."""
    from athena.ocr.contract import OCRBlock, OCRResult

    calls: list[dict[str, Any]] = []

    def _fn(image_path: Path, *, languages: list[str] | None = None) -> dict:
        calls.append({
            "image_path": str(image_path),
            "languages": list(languages) if languages else None,
        })
        return {
            "text": f"OCR'd text from {Path(image_path).name}",
            "blocks": [{
                "text": f"OCR'd text from {Path(image_path).name}",
                "bbox": [0, 0, 100, 30],
                "confidence": 88.0,
            }],
            "language": "eng",
        }

    _fn.calls = calls  # type: ignore[attr-defined]
    return _fn
