"""DOCX extractor via python-docx (T4-05.1).

python-docx parses Office Open XML directly — gives us body
paragraphs, heading style metadata, tables, and core
properties (title / author / created / modified) without a
binary dependency.

DOCX has no concept of pages at parse time (page layout is a
renderer concern); ``pages`` is reported as 0 and outline
entries carry ``page=0``. Tables get ``page=0`` for the same
reason. For paginated output a downstream tool would render
to PDF first; that's out of scope here.

Library isolation: every python-docx API call lives in this
file. The tool layer doesn't know it's there.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any

from ..result import (
    DocumentResult,
    OutlineEntry,
    TableData,
)

logger = logging.getLogger(__name__)


# Heading style name → outline level. python-docx exposes the
# style name as ``paragraph.style.name`` — "Heading 1" through
# "Heading 9" are the canonical ones; "Title" is treated as
# level 1 because that's how most documents use it.
_HEADING_PATTERN = re.compile(r"^Heading\s*(\d+)$", re.IGNORECASE)


def extract(path: Path | str) -> DocumentResult:
    """Parse a DOCX into a normalized DocumentResult."""
    from docx import Document  # python-docx; lazy import

    doc = Document(str(path))
    body_parts: list[str] = []
    outline: list[OutlineEntry] = []

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            body_parts.append(text)
        level = _heading_level(para)
        if level is not None and text:
            outline.append(OutlineEntry(
                level=level, title=text, page=0,
            ))

    tables = _extract_tables(doc)
    metadata = _extract_metadata(doc)

    full_text = "\n".join(body_parts)

    return DocumentResult(
        text=full_text,
        pages=0,  # DOCX has no page concept at parse time
        outline=outline,
        tables=tables,
        metadata=metadata,
        scanned_pages=[],
        figures=[],  # python-docx exposes images but not bbox;
                     # we leave figure extraction as a TODO for
                     # the format that benefits most (PDF)
        page_texts={},
    )


# ---------------------------------------------------------------
# internals
# ---------------------------------------------------------------


def _heading_level(paragraph: Any) -> int | None:
    """Map a paragraph's style name to an outline level, or
    None when the paragraph isn't a heading."""
    try:
        style_name = paragraph.style.name if paragraph.style else ""
    except Exception:  # noqa: BLE001
        return None
    if not style_name:
        return None
    name = style_name.strip()
    if name.lower() == "title":
        return 1
    match = _HEADING_PATTERN.match(name)
    if match:
        try:
            level = int(match.group(1))
            if 1 <= level <= 9:
                return level
        except ValueError:
            return None
    return None


def _extract_tables(doc: Any) -> list[TableData]:
    """python-docx tables have cells with text per row. Strip
    whitespace per cell; preserve ragged rows."""
    tables: list[TableData] = []
    for tbl in doc.tables:
        rows: list[list[str]] = []
        for row in tbl.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            tables.append(TableData(page=0, rows=rows))
    return tables


def _extract_metadata(doc: Any) -> dict[str, Any]:
    """Core properties — python-docx exposes these on
    ``doc.core_properties``."""
    try:
        props = doc.core_properties
    except Exception:  # noqa: BLE001
        return {"format": "DOCX"}

    def _stringy(v: Any) -> Any:
        if v is None:
            return None
        try:
            return v.isoformat()  # datetime
        except (AttributeError, TypeError):
            pass
        return str(v) if v else None

    return {
        "title": (props.title or "").strip() or None,
        "author": (props.author or "").strip() or None,
        "subject": (props.subject or "").strip() or None,
        "keywords": (props.keywords or "").strip() or None,
        "creator": (props.author or "").strip() or None,
        "creation_date": _stringy(getattr(props, "created", None)),
        "mod_date": _stringy(getattr(props, "modified", None)),
        "page_count": 0,
        "format": "DOCX",
    }
